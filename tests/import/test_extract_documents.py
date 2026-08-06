import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from tools.catalog_import.extract_documents import extract_document


class DocumentExtractionTests(unittest.TestCase):
    def test_extracts_paragraphs_tables_and_embedded_media(self):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            picture = root / 'fixture.png'
            Image.new('RGB', (32, 24), '#c52d23').save(picture)

            document = Document()
            document.add_heading('ППТС-18', level=1)
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Грузоподъёмность'
            table.cell(0, 1).text = '18 000 кг'
            table.cell(1, 0).text = 'Количество осей'
            table.cell(1, 1).text = '2'
            document.add_picture(str(picture))
            source = root / 'fixture.docx'
            document.save(source)

            extracted = extract_document(source, root / 'media')

            self.assertIn('ППТС-18', extracted['paragraphs'])
            self.assertIn(['Грузоподъёмность', '18 000 кг'], extracted['tables'][0])
            self.assertIn(['Количество осей', '2'], extracted['tables'][0])
            self.assertTrue(any(block['type'] == 'table' for block in extracted['blocks']))
            self.assertEqual(len(extracted['media']), 1)
            self.assertEqual(len(extracted['media'][0]['sha256']), 64)
            self.assertTrue((root / 'media' / extracted['media'][0]['filename']).is_file())


if __name__ == '__main__':
    unittest.main()
